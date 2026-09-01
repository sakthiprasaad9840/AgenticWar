"""Multi-format extraction and PHI/PII masking hard gate."""
from __future__ import annotations
import io, json, re, secrets
from pathlib import Path
from typing import Any
import pandas as pd
import hashlib
import uuid

PATTERNS = {
    "MEMBER": re.compile(r"\b(?:MBD|MBR)-[A-Za-z0-9_-]{6,20}\b", re.I),
    "SSN": re.compile(r"\b\d{3}-\d{2}-\d{4}\b"),
    "PHONE": re.compile(r"\b(?:\+?1[-. ]?)?(?:\(\d{3}\)|\d{3})[-. ]?\d{3}[-. ]?\d{4}\b"),
    "EMAIL": re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"),
    "MRN": re.compile(r"\bMRN[-_ ]?[A-Za-z0-9]{6,20}\b", re.I),
    "POLICY": re.compile(r"\bPOL[-_][A-Za-z0-9-]{6,30}\b", re.I),
}


"""
Member D — PHI/PII De-identification Utility (Enhanced)
Regex-based sanitizer to tokenize Member IDs, Claim IDs, TINs, SSNs, Emails, and Phone Numbers.
"""
def mask_phi_pii(text: str) -> tuple[str, dict[str, str]]:
    """
    Sanitizes patient-level PHI/PII while preserving operational business identifiers
    (TIN, CPT codes, Product, Claim IDs) necessary for database joins and rate resolution.
    """
    if not isinstance(text, str):
        return text, {}

    # Target strictly individual PHI/PII. 
    # Exclude operational keys (TIN, CPT, Product) to prevent pipeline join failures.
    patterns = {
        "EMAIL": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',
        # Require 10 digits so 9-digit TINs (XX-XXXXXXX) are not falsely captured
        "PHONE": r'\b(?:\+?\d{1,3}[-. ]?)?\(?\d{3}\)?[-. ]?\d{3}[-. ]?\d{4}\b',
        "SSN": r'\b\d{3}-\d{2}-\d{4}\b',
        "MEMBER": r'\b(?:MBD|MEM|PAT)-\d{6,9}\b',
        "IP_ADDR": r'\b(?:\d{1,3}\.){3}\d{1,3}\b'
    }

    spans = []
    # Collect all pattern match offsets
    for token_prefix, pattern in patterns.items():
        for match in re.finditer(pattern, text):
            spans.append((match.start(), match.end(), match.group(), token_prefix))

    # Sort spans descending by start index (right-to-left replacement prevents offset drift)
    spans.sort(key=lambda x: x[0], reverse=True)

    sanitized_text = text
    token_vault = {}
    
    last_start = len(text) + 1
    for start, end, match_text, token_prefix in spans:
        if end <= last_start:  # Guard against overlapping spans
            # Deterministic/Unique UUID slice prevents hash collision across separate runs
            short_id = uuid.uuid4().hex[:8].upper()
            token = f"[{token_prefix}_TOKEN_{short_id}]"
            
            token_vault[token] = match_text
            sanitized_text = sanitized_text[:start] + token + sanitized_text[end:]
            last_start = start

    return sanitized_text, token_vault


def mask_nested(value: Any) -> tuple[Any, dict[str, str]]:
    if isinstance(value, str):
        return mask_phi_pii(value)
    vault: dict[str, str] = {}
    if isinstance(value, list):
        out = []
        for item in value:
            masked, v = mask_nested(item); out.append(masked); vault.update(v)
        return out, vault
    if isinstance(value, tuple):
        out = []
        for item in value:
            masked, v = mask_nested(item); out.append(masked); vault.update(v)
        return tuple(out), vault
    if isinstance(value, dict):
        out = {}
        for key, item in value.items():
            masked_key, kv = mask_phi_pii(str(key)); masked_item, iv = mask_nested(item)
            out[masked_key] = masked_item; vault.update(kv); vault.update(iv)
        return out, vault
    return value, vault


def extract_docx_text(path: str | Path) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for i, table in enumerate(doc.tables, 1):
        parts.append(f"--- Table {i} ---")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells): parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx_text(path: str | Path) -> str:
    sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    parts = []
    for name, df in sheets.items():
        parts.append(f"--- Sheet: {name} ---")
        parts.append(df.fillna("").to_csv(index=False) if df is not None else "[EMPTY SHEET]")
    return "\n".join(parts)


def extract_csv_text(path: str | Path) -> str:
    return pd.read_csv(path, dtype=str, keep_default_na=False).to_csv(index=False)


def extract_json_text(path: str | Path) -> str:
    with open(path, encoding="utf-8") as f: data = json.load(f)
    return json.dumps(data, indent=2, ensure_ascii=False, default=str)


def extract_pdf_text(path: str | Path) -> str:
    import fitz
    doc = fitz.open(path); pages = []
    try:
        for n, page in enumerate(doc, 1):
            text = page.get_text("text").strip()
            if text:
                pages.append(f"--- Page {n} ---\n{text}"); continue
            try:
                import pytesseract
                from PIL import Image
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image = Image.open(io.BytesIO(pix.tobytes("png")))
                text = pytesseract.image_to_string(image).strip()
                pages.append(f"--- Page {n} OCR ---\n{text or '[NO TEXT EXTRACTED]'}")
            except Exception:
                pages.append(f"--- Page {n} ---\n[OCR UNAVAILABLE]")
    finally:
        doc.close()
    return "\n\n".join(pages)


def extract_document(path: str | Path) -> str:
    p = Path(path); ext = p.suffix.lower()
    if ext == ".pdf": return extract_pdf_text(p)
    if ext == ".docx": return extract_docx_text(p)
    if ext == ".xlsx": return extract_xlsx_text(p)
    if ext == ".csv": return extract_csv_text(p)
    if ext == ".json": return extract_json_text(p)
    if ext == ".txt": return p.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported file format: {ext}")


def extract_and_mask(path: str | Path) -> tuple[str, dict[str, str]]:
    return mask_phi_pii(extract_document(path))



def mask_document_file(path: str | Path, output_path: str | Path) -> tuple[Path, dict[str, str]]:
    """Create a masked copy of a PDF/DOCX contract and return its PHI/PII vault."""
    source = Path(path)
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    ext = source.suffix.lower()

    if ext == ".pdf":
        import fitz
        doc = fitz.open(source)
        vault: dict[str, str] = {}
        try:
            for page in doc:
                page_text = page.get_text("text") or ""
                masked_text, page_vault = mask_phi_pii(page_text)
                # Redact exact detected values so the original text is removed
                # from the PDF rather than merely hidden in the UI.
                for token, original in page_vault.items():
                    for rect in page.search_for(original):
                        page.add_redact_annot(rect, text=token)
                if page_vault:
                    page.apply_redactions()
                vault.update(page_vault)
            doc.save(target)
        finally:
            doc.close()
        return target, vault

    if ext == ".docx":
        from docx import Document
        doc = Document(source)
        vault: dict[str, str] = {}

        def mask_paragraph(paragraph):
            nonlocal vault
            original = paragraph.text
            masked, v = mask_phi_pii(original)
            if masked == original:
                return
            # Replacing paragraph text preserves the document structure while
            # ensuring the sensitive value is no longer present in the copy.
            paragraph.text = masked
            vault.update(v)

        for paragraph in doc.paragraphs:
            mask_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        mask_paragraph(paragraph)
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                mask_paragraph(paragraph)
            for paragraph in section.footer.paragraphs:
                mask_paragraph(paragraph)

        doc.save(target)
        return target, vault

    raise ValueError(f"Contract masking supports PDF and DOCX only, not {ext}.")
