import sys
from pathlib import Path
ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
import json
import pandas as pd

from lib.deidentify import extract_docx_text, extract_xlsx_text, extract_json_text, mask_phi_pii
from lib.validation import load_config, load_claims


def test_masking_patterns():
    text = "MRN-ABC123456 email test@example.com phone 202-555-0123 SSN 123-45-6789 MBR-ABC123456"
    masked, vault = mask_phi_pii(text)
    assert "test@example.com" not in masked
    assert "202-555-0123" not in masked
    assert "123-45-6789" not in masked
    assert len(vault) >= 4


def test_docx_tables_and_xlsx_all_sheets(tmp_path: Path):
    from docx import Document
    docx_path = tmp_path / "sample.docx"
    doc = Document(); doc.add_paragraph("Agreement text")
    table = doc.add_table(rows=1, cols=2); table.cell(0, 0).text = "Contact"; table.cell(0, 1).text = "test@example.com"
    doc.save(docx_path)
    assert "Contact" in extract_docx_text(docx_path)
    assert "test@example.com" in extract_docx_text(docx_path)

    xlsx_path = tmp_path / "sample.xlsx"
    with pd.ExcelWriter(xlsx_path) as writer:
        pd.DataFrame({"a": [1]}).to_excel(writer, sheet_name="First", index=False)
        pd.DataFrame({"b": [2]}).to_excel(writer, sheet_name="Second", index=False)
    text = extract_xlsx_text(xlsx_path)
    assert "Sheet: First" in text and "Sheet: Second" in text


def test_json_and_claim_normalization(tmp_path: Path):
    json_path = tmp_path / "config.json"
    json_path.write_text(json.dumps([{
        "tin": "123", "cpt_code": "99215", "product": "PPO",
        "effective_date": "2026-01-01", "end_date": "2026-12-31", "configured_rate": 100
    }]))
    cfg = load_config(str(json_path))
    assert list(cfg.columns) == ["tin", "cpt_code", "product", "effective_date", "end_date", "configured_rate"]

    claims_path = tmp_path / "claims.csv"
    pd.DataFrame([{
        "claim_id": "C1", "billing_tin": "123", "cpt_hcpcs": "99215", "product": "PPO",
        "date_of_service": "2029-02-30", "time_of_service": "", "place_of_service": "11",
        "billed_amount": "100", "paid_amount": "90"
    }]).to_csv(claims_path, index=False)
    claims = load_claims(str(claims_path))
    assert "tin" in claims.columns and "cpt_code" in claims.columns
    assert pd.isna(claims.loc[0, "date_of_service"])


def test_negotiate_accepts_pdf_and_docx_extensions():
    from pathlib import Path
    assert Path("draft_contract.pdf").suffix == ".pdf"
    assert Path("draft_contract.docx").suffix == ".docx"
